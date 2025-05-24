from funasr import AutoModel
import os
from docx import Document
import datetime
from markdown_it.rules_block import paragraph

# paraformer-zh is a multi-functional asr model
# use vad, punc, spk or not as you need
def video_to_text(path):
    model = AutoModel(model="paraformer-zh", model_revision="v2.0.4",
                      vad_model="fsmn-vad", vad_model_revision="v2.0.4",
                      punc_model="ct-punc-c", punc_model_revision="v2.0.4",
                      # spk_model="cam++", spk_model_revision="v2.0.2",
                      )

    # 获取当前脚本所在的目录
    path = path
    res = model.generate(input=path,
                batch_size_s=300)
    data = res[0]["text"]
    return data

def split_text_into_chunks_by_periods(text, periods_per_chunk):
    """将文本按指定数量的句号分割成多个段落"""
    chunks = []
    current_chunk = []
    period_count = 0

    # 将文本分割成句子
    sentences = text.split('。')  # 假设中文句号是段落的分隔符

    for sentence in sentences:
        current_chunk.append(sentence)
        period_count += 1

        # 当达到指定的句号数量时，开始新的一段
        if period_count == periods_per_chunk:
            # 将当前累积的句子合并为一个段落，并添加到chunks列表中
            chunks.append('。'.join(current_chunk) + '。')
            # 重置当前段落和句号计数
            current_chunk = []
            period_count = 0

    # 添加最后一部分，如果它包含少于指定数量的句号
    if current_chunk:
        chunks.append('。'.join(current_chunk) + ('。' if text.endswith('。') else ''))

    return chunks



# # 创建Word文档并添加段落
# from docx import Document
# doc = Document()
# for paratext in paragraphs:
#     doc.add_paragraph(paratext)
# doc.save('output.docx')

def save_to_word(video_path, periods_per_chunk,save_path):
    path = video_path
    paragraphs = split_text_into_chunks_by_periods(video_to_text(path),periods_per_chunk)
    doc = Document()
    for paratext in paragraphs:
        doc.add_paragraph(paratext)
    word_path = os.path.join(save_path,f"文本{datetime.datetime.today().date()}.docx")
    doc.save(word_path)
    return word_path

if __name__ == '__main__':
    pass
